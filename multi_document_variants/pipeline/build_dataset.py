import json
import docx
from pathlib import Path

ROOT_DIR = Path(__file__).resolve().parent.parent  # multi_document_variants/

PROFILE_DIR = ROOT_DIR / "data" / "profile_variants"

SYSTEM_PROMPT = (
    "Bạn là trợ lý tóm tắt văn bản cá nhân hóa cho cán bộ, công chức Việt Nam. "
    "Dựa vào văn bản gốc và thông tin cá nhân (persona) của người đọc, hãy viết một "
    "bản tóm tắt phù hợp: chọn lọc nội dung liên quan đến chủ đề quan tâm và lĩnh vực "
    "chuyên môn của người đọc, trình bày với mức độ chuyên sâu phù hợp, giữ đúng bố cục "
    "ưu tiên và giọng điệu, thái độ đúng đắn."
)

_persona_cache = {}


def get_persona_map(variant):
    ten_file = f"state_profiles_{variant}.json" if variant else "state_profiles_nt_nn_tc_kn_cd_ch.json"
    if ten_file not in _persona_cache:
        with open(PROFILE_DIR / ten_file, encoding="utf-8") as f:
            danh_sach = json.load(f)
        _persona_cache[ten_file] = {p["id"]: p for p in danh_sach}
    return _persona_cache[ten_file]


def dinh_dang_persona(persona):
    dong = [
        f"- Ngành/lĩnh vực: {persona.get('nganh_to', '')} / {persona.get('nganh_nho', '')}",
        f"- Tổ chức: {persona.get('to_chuc', '')}",
        f"- Kinh nghiệm: {persona.get('kinh_nghiem', '')}",
        f"- Chủ đề quan tâm (ưu tiên theo thứ tự): {', '.join(persona.get('chu_de', []) or [])}",
        f"- Xu hướng/mối quan tâm hiện tại: {persona.get('cau_hoi_truoc_mat', '')}",
        f"- Mô tả chung: {persona.get('mo_ta_chung', '')}",
    ]
    return "\n".join(dong)


def build_user_content(van_ban_goc_text, persona):
    return (
        f"THÔNG TIN NGƯỜI ĐỌC:\n{dinh_dang_persona(persona)}\n\n"
        f"VĂN BẢN GỐC:\n{van_ban_goc_text}\n\n"
        f"Hãy viết bản tóm tắt cá nhân hóa cho người đọc trên."
    )


def build_sample(nhanh, variant, nguon_goc, persona_id, van_ban_goc_text, tom_tat):
    persona = get_persona_map(variant).get(persona_id)
    if not persona:
        return None
    ten_variant_hien_thi = variant if variant else "full"
    return {
        "id": f"{nhanh}_{ten_variant_hien_thi}_{persona_id}_{nguon_goc}",
        "nhanh": nhanh,
        "variant": ten_variant_hien_thi,
        "persona_id": persona_id,
        "nguon_goc": nguon_goc,
        "messages": [
            {"role": "system", "content": SYSTEM_PROMPT},
            {"role": "user", "content": build_user_content(van_ban_goc_text, persona)},
            {"role": "assistant", "content": tom_tat},
        ],
    }


# ==== BÁO CHÍ ====

def eval_pass_chuan(eval_json):
    return eval_json.get("verdict_cuoi") == "DAT" and eval_json.get("so_tieu_chi_dat", 0) == 6


def _tach_variant_va_id(duong_dan_tuong_doi_parts):
    # 1 phan tu -> khong co variant; 2 phan tu -> [variant, id]
    if len(duong_dan_tuong_doi_parts) == 1:
        return None, duong_dan_tuong_doi_parts[0]
    return duong_dan_tuong_doi_parts[0], duong_dan_tuong_doi_parts[1]


def process_bao_chi():
    summary_root = ROOT_DIR / "output" / "bao_chi" / "rss_summary" / "json"
    eval_root = ROOT_DIR / "output" / "bao_chi" / "rss_summary" / "eval"
    snapshot_file = ROOT_DIR / "data" / "bao_chi" / "vnexpress_rss_snapshot_3007.json"

    with open(snapshot_file, "r", encoding="utf-8") as f:
        danh_sach_bai = json.load(f)
    van_ban_goc_text = "\n\n---\n\n".join(
        f"[{bai.get('category_name', '')}] {bai.get('title', '')}\n{bai.get('summary', '')}"
        for bai in danh_sach_bai
    )

    ket_qua = []
    bo_qua = 0
    for f_summary in summary_root.rglob("*.json"):
        parts = f_summary.relative_to(summary_root).with_suffix("").parts
        variant, persona_id = _tach_variant_va_id(parts)

        f_eval = eval_root / (Path(*parts[:-1]) if variant else Path()) / f"{persona_id}.json"
        if not f_eval.exists():
            bo_qua += 1
            continue
        with open(f_summary, "r", encoding="utf-8") as f:
            summary_json = json.load(f)
        with open(f_eval, "r", encoding="utf-8") as f:
            eval_json = json.load(f)
        if not eval_pass_chuan(eval_json):
            bo_qua += 1
            continue
        sample = build_sample(
            "bao_chi", variant, "vnexpress_rss_snapshot_3007", persona_id,
            van_ban_goc_text, summary_json["summary"]
        )
        if sample:
            ket_qua.append(sample)
        else:
            bo_qua += 1

    print(f"Báo chí: lấy được {len(ket_qua)} sample, bỏ qua {bo_qua}")
    return ket_qua


# ==== CHÍNH LUẬN ====

def eval_pass_chinh_luan(eval_json):
    return eval_json.get("dat_tat_ca_7_tieu_chi") is True


def process_chinh_luan():
    summary_root = ROOT_DIR / "output" / "chinh_luan" / "summary"
    eval_root = ROOT_DIR / "output" / "chinh_luan" / "eval"
    nguon_file = ROOT_DIR / "data" / "chinh_luan" / "nhandan_chinhluan.json"

    with open(nguon_file, "r", encoding="utf-8") as f:
        bai_map = {bai["id"]: bai for bai in json.load(f)}

    ket_qua = []
    bo_qua = 0
    for f_summary in summary_root.rglob("*.json"):
        parts = f_summary.relative_to(summary_root).with_suffix("").parts
        if len(parts) == 2:
            variant, bai_id, persona_id = None, parts[0], parts[1]
        else:
            variant, bai_id, persona_id = parts[0], parts[1], parts[2]

        f_eval = eval_root / (f"{variant}/{bai_id}" if variant else bai_id) / f"{persona_id}.json"
        if not f_eval.exists():
            bo_qua += 1
            continue
        with open(f_summary, "r", encoding="utf-8") as f:
            summary_json = json.load(f)
        with open(f_eval, "r", encoding="utf-8") as f:
            eval_json = json.load(f)
        if not eval_pass_chinh_luan(eval_json):
            bo_qua += 1
            continue

        bai_goc = bai_map.get(bai_id)
        if not bai_goc:
            bo_qua += 1
            continue
        van_ban_goc_text = f"{bai_goc['title']}\n\n{bai_goc['content']}"

        sample = build_sample("chinh_luan", variant, bai_id, persona_id, van_ban_goc_text, summary_json["summary"])
        if sample:
            ket_qua.append(sample)
        else:
            bo_qua += 1

    print(f"Chính luận: lấy được {len(ket_qua)} sample, bỏ qua {bo_qua}")
    return ket_qua


# ==== HÀNH CHÍNH ====

def trich_text_thuan(duong_dan_file):
    doc = docx.Document(duong_dan_file)
    dong = []
    for p in doc.paragraphs:
        if p.text.strip():
            dong.append(p.text.strip())
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    dong.append(cell.text.strip())
    return "\n".join(dong)


def resolve_docx_path(duong_dan_luu_trong_json):
    ten_file = duong_dan_luu_trong_json.replace("\\", "/").split("/")[-1]
    return ROOT_DIR / "data" / "hanh_chinh" / ten_file


def process_hanh_chinh():
    summary_root = ROOT_DIR / "output" / "hanh_chinh" / "summary"
    eval_root = ROOT_DIR / "output" / "hanh_chinh" / "eval"

    ket_qua = []
    bo_qua = 0
    docx_cache = {}
    for f_summary in summary_root.rglob("*.json"):
        parts = f_summary.relative_to(summary_root).with_suffix("").parts
        if len(parts) == 2:
            variant, ma_van_ban, persona_id = None, parts[0], parts[1]
        else:
            variant, ma_van_ban, persona_id = parts[0], parts[1], parts[2]

        f_eval = eval_root / (f"{variant}/{ma_van_ban}" if variant else ma_van_ban) / f"{persona_id}.json"
        if not f_eval.exists():
            bo_qua += 1
            continue
        with open(f_summary, "r", encoding="utf-8") as f:
            summary_json = json.load(f)
        with open(f_eval, "r", encoding="utf-8") as f:
            eval_json = json.load(f)
        if not eval_pass_chuan(eval_json):
            bo_qua += 1
            continue

        docx_path = resolve_docx_path(summary_json["file"])
        if not docx_path.exists():
            print(f"CẢNH BÁO: không tìm thấy file docx {docx_path}, bỏ qua {ma_van_ban}/{persona_id}")
            bo_qua += 1
            continue
        if docx_path not in docx_cache:
            docx_cache[docx_path] = trich_text_thuan(docx_path)
        van_ban_goc_text = docx_cache[docx_path]

        sample = build_sample("hanh_chinh", variant, ma_van_ban, persona_id, van_ban_goc_text, summary_json["summary"])
        if sample:
            ket_qua.append(sample)
        else:
            bo_qua += 1

    print(f"Hành chính: lấy được {len(ket_qua)} sample, bỏ qua {bo_qua}")
    return ket_qua

def main():
    theo_nhanh = {
        "bao_chi": process_bao_chi(),
        "chinh_luan": process_chinh_luan(),
        "hanh_chinh": process_hanh_chinh(),
    }

    output_dir = ROOT_DIR / "output" / "dataset"

    for nhanh, samples in theo_nhanh.items():
        theo_variant = {}
        for s in samples:
            theo_variant.setdefault(s["variant"], []).append(s)

        print(f"\n{nhanh}:")
        for variant, ds in sorted(theo_variant.items()):
            print(f"  - {variant}: {len(ds)} sample")

        nhanh_dir = output_dir / nhanh
        nhanh_dir.mkdir(parents=True, exist_ok=True)
        for variant, ds in theo_variant.items():
            f_out = nhanh_dir / f"{variant}.jsonl"
            with open(f_out, "w", encoding="utf-8") as f:
                for sample in ds:
                    f.write(json.dumps(sample, ensure_ascii=False) + "\n")

    tong = sum(len(v) for v in theo_nhanh.values())
    print(f"\nTổng cộng: {tong} sample, đã lưu tại {output_dir}")


if __name__ == "__main__":
    main()
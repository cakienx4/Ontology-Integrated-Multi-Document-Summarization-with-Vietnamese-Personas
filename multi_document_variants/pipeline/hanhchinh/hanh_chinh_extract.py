import re
import sys
import json
import docx
from pathlib import Path

ROOT_DIR = Path(__file__).resolve().parents[2]
# ==== CÁC HẰNG SỐ CẤU HÌNH ====

# từ khóa để nhận diện đoạn có khả năng liên quan tới đối tượng thi hành
TU_KHOA_DOI_TUONG = [
    "Kính gửi", "Nơi nhận", "chịu trách nhiệm thi hành",
    "yêu cầu", "Sở", "UBND", "Chủ tịch", "Chánh Văn phòng",
    "Giám đốc", "Trưởng ban", "Ban Quản lý"
]
# các cụm từ tương đương "chịu trách nhiệm thi hành" hay gặp trong Quyết định
CUM_TU_THI_HANH = [
    "chịu trách nhiệm thi hành",
    "có trách nhiệm thi hành",
    "tổ chức thi hành Quyết định",
    "thi hành Quyết định này"
]
TIEN_TO_TEN_CO_QUAN = [
    "Sở ", "UBND ", "Ủy ban nhân dân ", "Công an ", "Bộ Tư lệnh ",
    "Viện ", "Chi Cục ", "Chi cục ", "Văn phòng ", "Ban ", "Cơ quan ",
]
TU_KHOA_PHAN_CONG = ["chủ trì", "phối hợp"]
SO_KY_TU_TIM_TU_KHOA_PHAN_CONG = 80

def _la_doan_phan_cong(text):
    if any(text.startswith(tt) for tt in TIEN_TO_TEN_CO_QUAN):
        phan_can_kiem_tra = text
    else:
        vi_tri_giao = text.find("Giao ")
        if vi_tri_giao == -1:
            return False
        phan_sau_giao = text[vi_tri_giao + len("Giao "):]
        if not any(phan_sau_giao.startswith(tt) for tt in TIEN_TO_TEN_CO_QUAN):
            return False
        phan_can_kiem_tra = phan_sau_giao

    phan_dau_cau = phan_can_kiem_tra[:SO_KY_TU_TIM_TU_KHOA_PHAN_CONG]
    return any(tk in phan_dau_cau for tk in TU_KHOA_PHAN_CONG)

# số từ tối đa để coi 1 đoạn là "đoạn ngắn kiểu liệt kê" (dùng cho Kính gửi)
SO_TU_TOI_DA_DOAN_NGAN = 20

# số từ tối thiểu để loại các dòng bold ngắn kiểu "Nơi nhận:", "CHỦ TỊCH"
# ra khỏi danh sách heading giả
SO_TU_TOI_THIEU_HEADING_GIA = 3


# ==== BƯỚC 1: XÁC ĐỊNH LOẠI VĂN BẢN ====

def xac_dinh_loai_van_ban(doc):
    # ưu tiên 1: kiểm tra bảng quốc hiệu đầu tiên, chỉ khi đúng 1 dòng
    if len(doc.tables) > 0:
        bang_dau = doc.tables[0]
        if len(bang_dau.rows) == 1:
            text_bang = ""
            for row in bang_dau.rows:
                for cell in row.cells:
                    text_bang += cell.text.upper() + " "

            loai = _tim_tu_khoa_theo_kieu_chua(text_bang)
            if loai is not None:
                return loai

    # ưu tiên 2: quét danh sách heading, chỉ nhận heading BẮT ĐẦU BẰNG từ khóa
    danh_sach_heading = lay_danh_sach_heading(doc)
    for h in danh_sach_heading[:20]:
        text_upper = h["text"].strip().upper()
        loai = _tim_tu_khoa_theo_kieu_bat_dau(text_upper)
        if loai is not None:
            return loai

    # ưu tiên 2.5: một số văn bản để tên loại văn bản đứng riêng 1 dòng,
    # không in đậm nên không lọt vào danh sách heading (VD: KH-292 có
    # dòng "KẾ HOẠCH" đứng riêng, tách biệt khỏi tên kế hoạch phía dưới).
    # Chỉ quét đoạn RẤT NGẮN (dưới 15 ký tự) để chắc chắn đây là dòng
    # đứng riêng ghi tên loại văn bản, không phải câu trích dẫn.
    for p in doc.paragraphs[:15]:
        text = p.text.strip()
        if not text or len(text) > 15:
            continue
        loai = _tim_tu_khoa_theo_kieu_bat_dau(text.upper())
        if loai is not None:
            return loai

    # ưu tiên 3: riêng công văn, trích yếu "V/v..." thường không phải heading
    for p in doc.paragraphs[:15]:
        text = p.text.strip()
        if not text:
            continue
        if text.upper().startswith("V/V"):
            return "CONG_VAN"

    return "PHU_LUC_KHAC"


def _tim_tu_khoa_theo_kieu_chua(text_upper):
    if "CHỈ THỊ" in text_upper:
        return "CHI_THI"
    if "QUYẾT ĐỊNH" in text_upper:
        return "QUYET_DINH"
    if "KẾ HOẠCH" in text_upper:
        return "KE_HOACH"
    if "THÔNG BÁO" in text_upper:
        return "THONG_BAO"
    if "CÔNG VĂN" in text_upper:
        return "CONG_VAN"
    if "CÔNG ĐIỆN" in text_upper:
        return "CONG_DIEN"
    if "TỜ TRÌNH" in text_upper:
        return "TO_TRINH"
    if "BÁO CÁO" in text_upper:
        return "BAO_CAO"
    if "NGHỊ QUYẾT" in text_upper:
        return "NGHI_QUYET"
    return None


def _tim_tu_khoa_theo_kieu_bat_dau(text_upper):
    if text_upper.startswith("CHỈ THỊ"):
        return "CHI_THI"
    if text_upper.startswith("QUYẾT ĐỊNH"):
        return "QUYET_DINH"
    if text_upper.startswith("KẾ HOẠCH"):
        return "KE_HOACH"
    if text_upper.startswith("THÔNG BÁO"):
        return "THONG_BAO"
    if text_upper.startswith("BÁO CÁO"):
        return "BAO_CAO"
    if text_upper.startswith("CÔNG VĂN"):
        return "CONG_VAN"
    return None


# ==== BƯỚC 2: LẤY DANH SÁCH HEADING (CÓ FALLBACK) ====

def doan_la_heading(paragraph):
    """
    Kiểm tra 1 đoạn có phải heading không, có 3 trường hợp:
    1. Style tên có chữ "Heading" (heading chuẩn của Word)
    2. Style tên có chữ "Tiêu đề" (heading tự đặt tên, gặp ở văn bản dùng
       custom style như PL01)
    3. Cả đoạn in đậm (bold) nhưng style vẫn là Normal -> heading giả
    """
    ten_style = paragraph.style.name

    if "Heading" in ten_style:
        return True, "heading_chuan"

    if "Tiêu đề" in ten_style:
        return True, "heading_custom_style"

    # heading giả: toàn bộ run đều bold, và số từ đủ dài
    # (loại bỏ mấy dòng ngắn kiểu "Nơi nhận:", "CHỦ TỊCH")
    if paragraph.runs:
        tat_ca_bold = all(r.bold for r in paragraph.runs if r.text.strip())
        so_tu = len(paragraph.text.split())
        if tat_ca_bold and so_tu >= SO_TU_TOI_THIEU_HEADING_GIA:
            return True, "heading_gia_bold"

    return False, None


def lay_danh_sach_heading(doc):
    """
    Duyệt toàn bộ đoạn văn, trả về danh sách các đoạn được coi là heading,
    kèm theo vị trí (index) và loại heading để sau này dễ debug.
    """
    ket_qua = []
    for i, p in enumerate(doc.paragraphs):
        if not p.text.strip():
            continue
        la_heading, loai = doan_la_heading(p)
        if la_heading:
            ket_qua.append({
                "vi_tri": i,
                "text": p.text.strip(),
                "loai_heading": loai
            })
    return ket_qua


# ==== BƯỚC 3: TÁCH ĐOẠN NẾU 1 ĐOẠN CHỨA NHIỀU MỤC SỐ ====

def tach_doan_neu_gop_nhieu_muc(text):
    """
    Một số đoạn bị gộp nhiều mục số vào chung 1 paragraph, ví dụ:
    "5. Tiếp tục triển khai... 5.1. Trong quá trình..."
    Hàm này cố gắng tách ra thành list các mục con dựa theo pattern số thứ tự.

    Lưu ý: đây là xử lý best-effort bằng regex, không chắc đúng 100% với
    mọi trường hợp, nên sau khi tách cần người kiểm tra lại thủ công.
    """
    # pattern bắt các vị trí bắt đầu 1 mục số kiểu "5." hoặc "5.1."
    pattern = re.compile(r"(?<!\d)(\d{1,2}(?:\.\d{1,2})?\.)\s")
    vi_tri_bat_dau = [m.start() for m in pattern.finditer(text)]

    if len(vi_tri_bat_dau) <= 1:
        # không tách được hoặc chỉ có 1 mục, trả nguyên văn
        return [text]

    danh_sach_doan = []
    for i in range(len(vi_tri_bat_dau)):
        bat_dau = vi_tri_bat_dau[i]
        ket_thuc = vi_tri_bat_dau[i + 1] if i + 1 < len(vi_tri_bat_dau) else len(text)
        doan_con = text[bat_dau:ket_thuc].strip()
        if doan_con:
            danh_sach_doan.append(doan_con)

    return danh_sach_doan


# ==== BƯỚC 4: TRÍCH ĐỐI TƯỢNG THI HÀNH THEO TỪNG LOẠI VĂN BẢN ====

def trich_doi_tuong_chi_thi(doc):
    """
    Chỉ thị: đối tượng thi hành thường nằm ngay ở Heading 1 đầu văn bản,
    và rải rác trong các đoạn Normal có chứa từ khóa (Sở, UBND...).
    Không lấy lại đoạn đã lấy làm heading, và dừng quét khi gặp khối
    "Nơi nhận" (không phải đối tượng thi hành, chỉ là nơi lưu/gửi).
    """
    ket_qua = []
    text_da_lay = set()  # để dedupe, so sánh theo text đã chuẩn hóa
    danh_sach_heading = lay_danh_sach_heading(doc)

    # lấy heading 1 đầu tiên coi như đối tượng thi hành chính
    text_heading_dau_bai = None
    for h in danh_sach_heading:
        if h["loai_heading"] == "heading_chuan":
            text_heading_dau_bai = h["text"].strip()
            ket_qua.append({"nguon": "heading_dau_bai", "text": h["text"]})
            text_da_lay.add(text_heading_dau_bai)
            break

    # quét thêm các đoạn Normal có từ khóa liên quan
    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue

        # dừng quét khi gặp khối "Nơi nhận" - phần còn lại của văn bản
        # (nơi nhận, chữ ký...) không phải đối tượng thi hành
        if text.startswith("Nơi nhận"):
            break

        # bỏ qua nếu trùng với heading đã lấy ở trên
        if text in text_da_lay:
            continue

        if any(tk in text for tk in TU_KHOA_DOI_TUONG):
            # có thể đoạn này gộp nhiều mục số, thử tách ra trước
            cac_doan_con = tach_doan_neu_gop_nhieu_muc(text)
            for doan_con in cac_doan_con:
                doan_con_strip = doan_con.strip()
                if doan_con_strip in text_da_lay:
                    continue
                ket_qua.append({"nguon": "doan_noi_dung", "text": doan_con})
                text_da_lay.add(doan_con_strip)

    return ket_qua

def trich_doi_tuong_thong_bao(doc):
    ket_qua = []
    danh_sach_doan = doc.paragraphs

    # trích yếu: đoạn bắt đầu bằng "Về việc" trong 15 đoạn đầu
    for p in danh_sach_doan[:15]:
        text = p.text.strip()
        if not text:
            continue
        if text.upper().startswith("VỀ VIỆC"):
            ket_qua.append({"nguon": "trich_yeu", "text": text})
            break

    # "Kính gửi:" - lấy nguyên cả đoạn vì đối tượng nằm chung dòng này
    for p in danh_sach_doan:
        text = p.text.strip()
        if not text:
            continue
        if "Kính gửi" in text:
            ket_qua.append({"nguon": "kinh_gui", "text": text})
            break

    # quét phần nội dung còn lại, dừng khi gặp khối "Nơi nhận"
    for p in danh_sach_doan:
        text = p.text.strip()
        if not text:
            continue
        if text.startswith("Nơi nhận"):
            break

        if _la_doan_phan_cong(text):
            ket_qua.append({"nguon": "doan_phan_cong", "text": text})
        elif "yêu cầu" in text.lower() and ("Nhà đầu tư" in text or "Sở" in text):
            cac_doan_con = tach_doan_neu_gop_nhieu_muc(text)
            for doan_con in cac_doan_con:
                ket_qua.append({"nguon": "doan_ket", "text": doan_con})

    return ket_qua

def trich_doi_tuong_cong_van(doc):
    ket_qua = []
    danh_sach_doan = doc.paragraphs

    # tìm trích yếu "V/v..." trong 15 đoạn đầu, thêm vào trước tiên
    # để LLM có ngữ cảnh chủ đề công văn khi khớp ngành
    for p in danh_sach_doan[:15]:
        text = p.text.strip()
        if not text:
            continue
        if text.upper().startswith("V/V"):
            ket_qua.append({"nguon": "trich_yeu", "text": text})
            break

    vi_tri_kinh_gui = None
    for i, p in enumerate(danh_sach_doan):
        if "Kính gửi" in p.text:
            vi_tri_kinh_gui = i
            break

    if vi_tri_kinh_gui is None:
        # không tìm thấy Kính gửi, coi như văn bản không theo mẫu chuẩn
        return ket_qua

    # gom các đoạn ngắn ngay sau "Kính gửi:" cho đến khi gặp đoạn dài
    # (đoạn dài là dấu hiệu đã sang phần nội dung chính)
    for p in danh_sach_doan[vi_tri_kinh_gui + 1:]:
        text = p.text.strip()
        if not text:
            continue
        so_tu = len(text.split())
        if so_tu > SO_TU_TOI_DA_DOAN_NGAN:
            break
        ket_qua.append({"nguon": "kinh_gui", "text": text})

    return ket_qua

def trich_doi_tuong_phan_cong_trong_doan(doc):
    ket_qua = []
    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue
        if _la_doan_phan_cong(text):
            ket_qua.append({"nguon": "doan_phan_cong", "text": text})
    return ket_qua

def _chuan_hoa_khoang_trang(text):
    return " ".join(text.split())

def trich_doi_tuong_bao_cao(doc):
    ket_qua = []
    danh_sach_doan = doc.paragraphs

    for p in danh_sach_doan[:15]:
        text_chuan_hoa = _chuan_hoa_khoang_trang(p.text)
        if not text_chuan_hoa:
            continue
        if text_chuan_hoa.upper().startswith("BÁO CÁO"):
            tieu_de = text_chuan_hoa[len("BÁO CÁO"):].strip()
            if tieu_de:
                ket_qua.append({"nguon": "trich_yeu", "text": tieu_de})
                ket_qua.append({"nguon": "chu_de_toan_van_ban", "text": tieu_de})
            break

    return ket_qua

def trich_doi_tuong_ke_hoach(doc):
    ket_qua = []
    danh_sach_heading = lay_danh_sach_heading(doc)

    da_lay_trich_yeu = False
    for h in danh_sach_heading[:10]:
        if h["text"].strip().upper().startswith("KẾ HOẠCH"):
            ket_qua.append({"nguon": "trich_yeu", "text": h["text"]})
            da_lay_trich_yeu = True
            break

    # fallback: một số văn bản (VD: KH-292) để tên loại văn bản "KẾ HOẠCH"
    # đứng riêng 1 dòng, không in đậm nên không lọt vào danh_sach_heading.
    # Trong trường hợp đó, heading_chuan đầu tiên sau bảng quốc hiệu
    # (vi_tri != 0) chính là tên kế hoạch thật, lấy làm trich_yeu.
    if not da_lay_trich_yeu:
        for h in danh_sach_heading[:10]:
            if h["vi_tri"] == 0:
                continue
            if h["loai_heading"] == "heading_chuan":
                ket_qua.append({"nguon": "trich_yeu", "text": h["text"]})
                da_lay_trich_yeu = True
                break

    for h in danh_sach_heading:
        if h["loai_heading"] == "heading_chuan" and "Giao" in h["text"]:
            ket_qua.append({"nguon": "heading_to_chuc_thuc_hien", "text": h["text"]})

    # câu phân công nhiệm vụ cụ thể theo từng cơ quan, nằm rải rác trong
    # đoạn văn thường (vd "Sở Y tế chủ trì phát triển y tế cơ sở.")
    ket_qua.extend(trich_doi_tuong_phan_cong_trong_doan(doc))

    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue
        if "yêu cầu" in text.lower() and "Sở" in text:
            cac_doan_con = tach_doan_neu_gop_nhieu_muc(text)
            for doan_con in cac_doan_con:
                ket_qua.append({"nguon": "doan_ket", "text": doan_con})

    return ket_qua


def trich_doi_tuong_quyet_dinh(doc):
    ket_qua = []

    # trích yếu dạng 1: dòng riêng "Về việc..." ngay sau tiêu đề "QUYẾT ĐỊNH"
    # (giống cấu trúc Thông báo)
    da_lay_trich_yeu = False
    for p in doc.paragraphs[:15]:
        text = p.text.strip()
        if not text:
            continue
        if text.upper().startswith("VỀ VIỆC"):
            ket_qua.append({"nguon": "trich_yeu", "text": text})
            da_lay_trich_yeu = True
            break

    # trích yếu dạng 2: gộp chung trong heading "QUYẾT ĐỊNH Về việc..."
    # (giống cấu trúc Báo cáo/Kế hoạch)
    if not da_lay_trich_yeu:
        danh_sach_heading = lay_danh_sach_heading(doc)
        for h in danh_sach_heading[:10]:
            text_chuan_hoa = _chuan_hoa_khoang_trang(h["text"])
            if text_chuan_hoa.upper().startswith("QUYẾT ĐỊNH"):
                tieu_de = text_chuan_hoa[len("QUYẾT ĐỊNH"):].strip(" :")
                if tieu_de:
                    ket_qua.append({"nguon": "trich_yeu", "text": tieu_de})
                break

    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue
        if text.startswith("Nơi nhận"):
            break
        if any(cum in text for cum in CUM_TU_THI_HANH):
            ket_qua.append({"nguon": "dieu_thi_hanh", "text": text})

    return ket_qua


def trich_doi_tuong_phu_luc(doc):
    ket_qua = []
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if not text:
            continue
        if "Phạm vi điều chỉnh" in text or "Đối tượng áp dụng" in text:
            # lấy luôn 2 đoạn ngay sau để có thêm ngữ cảnh
            doan_ke_tiep = doc.paragraphs[i + 1:i + 3]
            text_them = " ".join(pp.text.strip() for pp in doan_ke_tiep if pp.text.strip())
            ket_qua.append({
                "nguon": "chuong_quy_dinh_chung",
                "text": text + " " + text_them
            })

    if not ket_qua:
        ket_qua.append({
            "nguon": "khong_xac_dinh_duoc",
            "text": "Không tìm thấy Kính gửi / Nơi nhận / Phạm vi điều chỉnh, "
                    "cần kiểm tra thủ công."
        })

    return ket_qua


def trich_doi_tuong_thi_hanh(doc, loai_van_ban):
    if loai_van_ban == "CHI_THI":
        return trich_doi_tuong_chi_thi(doc)
    if loai_van_ban == "CONG_VAN":
        return trich_doi_tuong_cong_van(doc)
    if loai_van_ban == "KE_HOACH":
        return trich_doi_tuong_ke_hoach(doc)
    if loai_van_ban == "QUYET_DINH":
        return trich_doi_tuong_quyet_dinh(doc)
    if loai_van_ban == "THONG_BAO":
        return trich_doi_tuong_thong_bao(doc)
    if loai_van_ban == "BAO_CAO":
        return trich_doi_tuong_bao_cao(doc)

    return trich_doi_tuong_phu_luc(doc)


# ==== BƯỚC 5: TRÍCH BẢNG (CÓ XỬ LÝ MERGE Ô ĐẦU BẢNG) ====

def trich_bang(doc):
    ket_qua = []
    for ti, table in enumerate(doc.tables):
        so_dong = len(table.rows)
        so_cot = len(table.columns)

        header_row = []
        if so_dong > 0:
            hang_dau = table.rows[0].cells
            da_them = set()
            for cell in hang_dau:
                # cell._tc là ô XML thật, nếu 2 cell cùng trỏ 1 ô thật
                # nghĩa là bị merge, chỉ lấy 1 lần
                dinh_danh_o = id(cell._tc)
                if dinh_danh_o in da_them:
                    continue
                da_them.add(dinh_danh_o)
                header_row.append(cell.text.strip())

        ket_qua.append({
            "bang_so": ti,
            "so_dong": so_dong,
            "so_cot": so_cot,
            "header": header_row
        })

    return ket_qua

# ==== BƯỚC 5.5: TRÍCH TOÀN BỘ NỘI DUNG VĂN BẢN THEO MỤC (PHỤC VỤ TÓM TẮT) ====

def trich_toan_bo_noi_dung_theo_muc(doc):
    danh_sach_muc = []
    muc_hien_tai = {"heading": None, "vi_tri_heading": None, "doan_van": []}

    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if not text:
            continue

        if text.startswith("Nơi nhận"):
            break

        la_heading, loai_heading = doan_la_heading(p)
        if la_heading:
            if muc_hien_tai["heading"] is not None or muc_hien_tai["doan_van"]:
                danh_sach_muc.append(muc_hien_tai)
            muc_hien_tai = {"heading": text, "vi_tri_heading": i, "doan_van": []}
        else:
            muc_hien_tai["doan_van"].append(text)

    if muc_hien_tai["heading"] is not None or muc_hien_tai["doan_van"]:
        danh_sach_muc.append(muc_hien_tai)

    return danh_sach_muc


def _chuan_hoa_text_de_so_khop(text):

    return " ".join(text.split()).lower()

def tach_muc_theo_phan_cong_co_quan(danh_sach_muc):
    danh_sach_muc_moi = []

    for muc in danh_sach_muc:
        doan_van = muc.get("doan_van", [])

        co_doan_phan_cong = any(_la_doan_phan_cong(dv) for dv in doan_van)
        if not co_doan_phan_cong:
            danh_sach_muc_moi.append(muc)
            continue

        heading_goc = muc.get("heading")
        vi_tri_goc = muc.get("vi_tri_heading")
        da_dung_heading_goc = False
        cum_doan_thuong = []

        for dv in doan_van:
            if _la_doan_phan_cong(dv):
                if cum_doan_thuong:
                    danh_sach_muc_moi.append({
                        "heading": heading_goc if not da_dung_heading_goc else None,
                        "vi_tri_heading": vi_tri_goc,
                        "doan_van": cum_doan_thuong,
                    })
                    da_dung_heading_goc = True
                    cum_doan_thuong = []

                danh_sach_muc_moi.append({
                    "heading": heading_goc if not da_dung_heading_goc else None,
                    "vi_tri_heading": vi_tri_goc,
                    "doan_van": [dv],
                })
                da_dung_heading_goc = True
            else:
                cum_doan_thuong.append(dv)

        if cum_doan_thuong:
            danh_sach_muc_moi.append({
                "heading": heading_goc if not da_dung_heading_goc else None,
                "vi_tri_heading": vi_tri_goc,
                "doan_van": cum_doan_thuong,
            })

    return danh_sach_muc_moi

def danh_dau_muc_lien_quan(danh_sach_muc, doi_tuong_thi_hanh):
    for muc in danh_sach_muc:
        muc["chi_so_doi_tuong_lien_quan"] = []

        text_muc_gop = _chuan_hoa_text_de_so_khop(
            (muc["heading"] or "") + " " + " ".join(muc["doan_van"])
        )

        for idx_dt, dt in enumerate(doi_tuong_thi_hanh):
            if dt.get("nguon") == "chu_de_toan_van_ban":
                # đối tượng có phạm vi TOÀN VĂN BẢN (vd chủ đề Báo cáo) -
                # áp dụng cho MỌI mục, không cần so khớp substring
                muc["chi_so_doi_tuong_lien_quan"].append(idx_dt)
                continue

            text_dt = _chuan_hoa_text_de_so_khop(dt["text"])
            phan_dau = text_dt[:60]
            if phan_dau and phan_dau in text_muc_gop:
                muc["chi_so_doi_tuong_lien_quan"].append(idx_dt)

    return danh_sach_muc

# ==== BƯỚC 6: HÀM CHẠY CHÍNH CHO 1 FILE ====

def xu_ly_1_file(duong_dan_file):
    doc = docx.Document(duong_dan_file)

    loai_van_ban = xac_dinh_loai_van_ban(doc)
    danh_sach_heading = lay_danh_sach_heading(doc)
    doi_tuong_thi_hanh = trich_doi_tuong_thi_hanh(doc, loai_van_ban)
    danh_sach_bang = trich_bang(doc)
    CAC_LOAI_DA_CO_HAM_RIENG = {"CHI_THI", "CONG_VAN", "KE_HOACH", "QUYET_DINH", "THONG_BAO", "BAO_CAO"}
    if loai_van_ban not in CAC_LOAI_DA_CO_HAM_RIENG:
        print(f"CẢNH BÁO: {duong_dan_file} thuộc loại '{loai_van_ban}' chưa có "
              f"hàm trích riêng - đang dùng fallback chung, cần kiểm tra thủ công.")
    danh_sach_muc = trich_toan_bo_noi_dung_theo_muc(doc)
    danh_sach_muc = tach_muc_theo_phan_cong_co_quan(danh_sach_muc)
    danh_sach_muc = danh_dau_muc_lien_quan(danh_sach_muc, doi_tuong_thi_hanh)

    ket_qua = {
        "file": duong_dan_file,
        "loai_van_ban": loai_van_ban,
        "so_luong_heading": len(danh_sach_heading),
        "danh_sach_heading": danh_sach_heading,
        "doi_tuong_thi_hanh": doi_tuong_thi_hanh,
        "danh_sach_bang": danh_sach_bang,
        "danh_sach_muc": danh_sach_muc,
    }

    return ket_qua


def in_ket_qua_ra_man_hinh(ket_qua):
    """
    In kết quả ra màn hình dạng dễ đọc, không cần mở file JSON.
    """
    print(f"\n========== {ket_qua['file']} ==========")
    print(f"Loại văn bản xác định: {ket_qua['loai_van_ban']}")

    print(f"\n--- Heading ({ket_qua['so_luong_heading']}) ---")
    for h in ket_qua["danh_sach_heading"]:
        print(f"  [{h['vi_tri']}] ({h['loai_heading']}) {h['text'][:70]}")

    print(f"\n--- Đối tượng thi hành ({len(ket_qua['doi_tuong_thi_hanh'])}) ---")
    for dt in ket_qua["doi_tuong_thi_hanh"]:
        print(f"  ({dt['nguon']}) {dt['text'][:100]}")

    print(f"\n--- Bảng ({len(ket_qua['danh_sach_bang'])}) ---")
    for b in ket_qua["danh_sach_bang"]:
        print(f"  Bảng {b['bang_so']}: {b['so_dong']}x{b['so_cot']} header={b['header']}")


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Cách dùng: python -m pipeline.rss.hanhchinh.hanh_chinh_extract <đường_dẫn_file_docx>")
        sys.exit(1)

    duong_dan = Path(sys.argv[1])

    if not duong_dan.exists():
        duong_dan = ROOT_DIR / "data" / "hanh_chinh" / duong_dan

    ket_qua = xu_ly_1_file(str(duong_dan))
    in_ket_qua_ra_man_hinh(ket_qua)

    ten_file_json = duong_dan.with_name(f"{duong_dan.stem}_ket_qua.json")
    with open(ten_file_json, "w", encoding="utf-8") as f:
        json.dump(ket_qua, f, ensure_ascii=False, indent=2)

    print(f"\nĐã lưu kết quả JSON tại: {ten_file_json}")